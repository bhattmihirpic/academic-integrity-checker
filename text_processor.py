# text_processor.py

import os
import re
import io
from docx import Document  # For .docx files
from pdfminer.high_level import extract_text as pdf_extract_text  # Primary PDF extraction
from PyPDF2 import PdfReader  # Secondary PDF extraction
from pdf2image import convert_from_bytes  # For converting PDF pages to images
from PIL import Image  # For OCR image handling
import pytesseract  # For OCR
from werkzeug.datastructures import FileStorage

def extract_text_from_file_storage(file: FileStorage) -> str:
    """
    Extract text from a FileStorage object based on its extension.
    """
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == '.pdf':
        return read_pdf_from_storage(file)
    elif ext == '.docx':
        return read_word_from_storage(file)
    elif ext in ['.txt', '.text']:
        return read_text_from_storage(file)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def read_pdf_from_storage(file: FileStorage) -> str:
    """
    Robust PDF text extraction from FileStorage:
      1. Try pdfminer.six
      2. Fallback to PyPDF2
    """
    text = ""
    file_content = file.read()
    file.seek(0)  # Reset file pointer for potential reuse
    
    # 1. pdfminer.six extraction
    try:
        text = pdf_extract_text(io.BytesIO(file_content))
        print(f"[PDFMINER] Extracted {len(text)} chars from {file.filename}")
    except Exception as e:
        print(f"[PDFMINER ERROR] {e}")

    # 2. PyPDF2 fallback
    if not text.strip():
        try:
            reader = PdfReader(io.BytesIO(file_content))
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text())
            text = "\n".join(pages)
            print(f"[PYPDF2] Extracted {len(text)} chars from {file.filename}")
        except Exception as e:
            print(f"[PYPDF2 ERROR] {e}")

    return text

def read_word_from_storage(file: FileStorage) -> str:
    """Read text from a .docx file stored in FileStorage"""
    doc = Document(io.BytesIO(file.read()))
    file.seek(0)  # Reset file pointer
    
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def read_text_from_storage(file: FileStorage) -> str:
    """Read text from a .txt file stored in FileStorage"""
    return file.read().decode('utf-8')

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file on disk based on its extension.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return read_pdf_file(file_path)
    elif ext == '.docx':
        return read_word_file(file_path)
    elif ext in ['.txt', '.text']:
        return read_text_file(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def read_pdf_file(file_path: str) -> str:
    """
    Read text from a PDF file on disk
    """
    text = ""
    try:
        text = pdf_extract_text(file_path)
    except Exception as e:
        print(f"[PDFMINER ERROR] {e}")
        try:
            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text())
            text = "\n".join(pages)
        except Exception as e:
            print(f"[PYPDF2 ERROR] {e}")
    return text

def read_word_file(file_path: str) -> str:
    """Read text from a .docx file on disk"""
    doc = Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def read_text_file(file_path: str) -> str:
    """Read text from a .txt file on disk"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def clean_up_text(text: str) -> str:
    """
    Clean up extracted text by:
    - Removing extra whitespace
    - Converting to lowercase
    - Removing special characters
    """
    if not text:
        return ""
    
    # Convert to lowercase
    text = text.lower()
    
    # Replace multiple spaces with single space
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep periods and basic punctuation
    text = re.sub(r'[^a-z0-9\s.,!?]', '', text)
    
    # Remove extra whitespace
    text = text.strip()
    
    return text

def get_text_info(text: str) -> dict:
    """
    Get comprehensive text analysis including:
    - Basic stats (words, chars, sentences)
    - Readability metrics
    - Text complexity analysis
    - Vocabulary richness
    """
    if not text:
        return {
            "total_words": 0,
            "total_characters": 0,
            "total_sentences": 0,
            "avg_word_length": 0,
            "avg_sentence_length": 0,
            "unique_words": 0,
            "vocabulary_richness": 0,
            "readability_score": 0
        }
    
    # Basic stats
    words = text.split()
    total_words = len(words)
    total_chars = len(text)
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    total_sentences = len(sentences)
    
    # Average lengths
    avg_word_length = sum(len(word) for word in words) / total_words if total_words > 0 else 0
    avg_sentence_length = total_words / total_sentences if total_sentences > 0 else 0
    
    # Vocabulary richness
    unique_words = len(set(word.lower() for word in words))
    vocabulary_richness = (unique_words / total_words * 100) if total_words > 0 else 0
    
    # Simple readability score (based on avg sentence length and avg word length)
    # Lower score means easier to read (scale 0-100)
    readability_score = min(100, (avg_sentence_length * 0.5 + avg_word_length * 10))
    
    return {
        "total_words": total_words,
        "total_characters": total_chars,
        "total_sentences": total_sentences,
        "avg_word_length": round(avg_word_length, 1),
        "avg_sentence_length": round(avg_sentence_length, 1),
        "unique_words": unique_words,
        "vocabulary_richness": round(vocabulary_richness, 1),
        "readability_score": round(readability_score, 1)
    }
